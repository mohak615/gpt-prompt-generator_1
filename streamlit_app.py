import asyncio
import sys
if sys.platform.startswith('win'):
    asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

import streamlit as st
import json
import re
import time
import os
import openai
from scraper import scrape_business_info_with_ai
from docx import Document
from docx.shared import Inches
import io
import base64

# Page configuration
st.set_page_config(
    page_title="GPT Prompt Generator",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        text-align: center;
        color: #1f77b4;
        font-size: 2.5rem;
        font-weight: bold;
        margin-bottom: 2rem;
    }
    .sub-header {
        color: #ff7f0e;
        font-size: 1.2rem;
        margin-bottom: 1rem;
    }
    .success-box {
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        border-radius: 5px;
        padding: 1rem;
        margin: 1rem 0;
    }
    .error-box {
        background-color: #f8d7da;
        border: 1px solid #f5c6cb;
        border-radius: 5px;
        padding: 1rem;
        margin: 1rem 0;
    }
    .info-box {
        background-color: #d1ecf1;
        border: 1px solid #bee5eb;
        border-radius: 5px;
        padding: 1rem;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# OpenAI API Key
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

def load_template(template_file="prompt_template.txt"):
    """Load the prompt template from file"""
    try:
        with open(template_file, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        st.error(f"❌ Error: Template file '{template_file}' not found.")
        return None
    except Exception as e:
        st.error(f"❌ Error reading template file: {e}")
        return None

def map_scraped_data_to_template(scraped_data):
    """Map scraped data fields to template placeholders"""
    mapping = {
        # Basic Info
        "{{company_name}}": scraped_data.get("company_name", "Not available"),
        "{{address}}": scraped_data.get("address", "Not available"),
        "{{phone_number}}": scraped_data.get("phone_number", "Not available"),
        "{{email}}": scraped_data.get("email", "Not available"),
        "{{website}}": scraped_data.get("website_url", "Not available"),
        "{{business_hours}}": scraped_data.get("business_hours", "Not available"),
        "{{timezone}}": scraped_data.get("timezone", "Not available"),
        
        # Services
        "{{services}}": scraped_data.get("services_list", "Not available"),
        "{{service_descriptions}}": scraped_data.get("service_descriptions", "Not available"),
        "{{pricing}}": scraped_data.get("pricing", "Not available"),
        "{{duration}}": scraped_data.get("duration", "Not available"),
        "{{booking_links}}": scraped_data.get("booking_links", "Not available"),
        "{{service_areas}}": scraped_data.get("service_areas", "Not available"),
        
        # Payments & Policies
        "{{payment_methods}}": scraped_data.get("payment_methods", "Not available"),
        "{{financing}}": scraped_data.get("financing_plans", "Not available"),
        "{{refund_policy}}": scraped_data.get("refund_policy", "Not available"),
        
        # Team Information
        "{{team}}": scraped_data.get("staff_names", "Not available"),
        "{{bios}}": scraped_data.get("staff_titles", "Not available"),
        "{{team_photos}}": scraped_data.get("staff_photos", "Not available"),
        
        # Social Media
        "{{facebook}}": scraped_data.get("facebook_url", "Not available"),
        "{{instagram}}": scraped_data.get("instagram_url", "Not available"),
        "{{linkedin}}": scraped_data.get("linkedin_url", "Not available"),
        "{{other_socials}}": scraped_data.get("social_handles", "Not available"),
        
        # Reviews / Testimonials
        "{{testimonials}}": scraped_data.get("testimonials", "Not available"),
        
        # Policies & Legal
        "{{privacy_policy}}": scraped_data.get("privacy_policy", "Not available"),
        "{{refund_policy}}": scraped_data.get("refund_policy", "Not available"),
        "{{terms}}": scraped_data.get("terms_of_service", "Not available"),
        "{{licenses}}": scraped_data.get("licenses_certifications", "Not available"),
        
        # Brand & Philosophy
        "{{tagline}}": scraped_data.get("tagline", "Not available"),
        "{{mission_statement}}": scraped_data.get("mission_statement", "Not available"),
        "{{tone}}": scraped_data.get("communication_style", "Not available"),
    }
    return mapping

def replace_placeholders(template, data_mapping):
    """Replace all placeholders in the template with actual data"""
    result = template
    for placeholder, value in data_mapping.items():
        result = result.replace(placeholder, str(value))
    return result

def test_gpt_response(final_prompt, model="gpt-3.5-turbo"):
    """Test GPT API with the generated prompt"""
    try:
        # Configure OpenAI
        openai.api_key = OPENAI_API_KEY
        
        # Create a test prompt
        test_prompt = f"""
Based on the following business information, provide a comprehensive analysis and suggestions:

{final_prompt}

Please provide:
1. Business Overview (2-3 sentences)
2. Key Strengths (3-4 points)
3. Marketing Suggestions (3-4 ideas)
4. Potential Improvements (2-3 suggestions)
5. Target Audience Analysis (1-2 sentences)

Format your response professionally with clear sections.
"""
        
        # Make API call
        response = openai.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are a professional business analyst and marketing consultant. Provide clear, actionable insights based on business information."},
                {"role": "user", "content": test_prompt}
            ],
            max_tokens=800,
            temperature=0.7
        )
        
        gpt_response = response.choices[0].message.content
        
        # Show response stats
        response_tokens = response.usage.total_tokens
        prompt_tokens = response.usage.prompt_tokens
        completion_tokens = response.usage.completion_tokens
        
        return True, gpt_response, {
            'total_tokens': response_tokens,
            'prompt_tokens': prompt_tokens,
            'completion_tokens': completion_tokens
        }
        
    except Exception as e:
        return False, str(e), None

def create_docx_download(final_prompt, gpt_response, company_name="Business"):
    """Create a DOCX file for download"""
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Business Analysis Report - {company_name}', 0)
    title.alignment = 1  # Center alignment
    
    # Add timestamp
    from datetime import datetime
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f"Generated on: {timestamp}")
    doc.add_paragraph("")  # Empty line
    
    # Add original prompt section
    doc.add_heading('Original Business Information', level=1)
    doc.add_paragraph(final_prompt)
    doc.add_paragraph("")  # Empty line
    
    # Add GPT response section
    if gpt_response:
        doc.add_heading('AI Analysis & Recommendations', level=1)
        doc.add_paragraph(gpt_response)
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes

def get_download_link(file_bytes, filename, text):
    """Generate download link for file"""
    b64 = base64.b64encode(file_bytes.read()).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">{text}</a>'
    return href

def main():
    # Header
    st.markdown('<h1 class="main-header">🚀 GPT Prompt Generator</h1>', unsafe_allow_html=True)
    st.markdown('<p style="text-align: center; font-size: 1.1rem; color: #666;">Professional Business Data Scraper & Prompt Builder</p>', unsafe_allow_html=True)
    
    # Sidebar
    with st.sidebar:
        st.header("⚙️ Settings")
        st.markdown("### Configuration")
        
        # Model selection
        model = st.selectbox(
            "GPT Model",
            ["gpt-3.5-turbo", "gpt-4", "gpt-4-turbo"],
            index=0
        )
        
        # Max pages for scraping
        max_pages = st.slider("Max Pages to Scrape", 3, 15, 8)
        
        # Test GPT API option
        test_gpt = st.checkbox("Test with GPT API", value=True)
        
        st.markdown("---")
        st.markdown("### About")
        st.markdown("""
        This tool extracts comprehensive business information from websites and generates professional AI prompts.
        
        **Features:**
        - 🔍 Web scraping with AI
        - 📝 Template-based prompt generation
        - 🤖 GPT API integration
        - 📄 DOCX export
        """)
    
    # Main content
    st.markdown('<h2 class="sub-header">🌐 Enter Business Website URL</h2>', unsafe_allow_html=True)
    
    # URL input
    url = st.text_input(
        "Website URL",
        placeholder="https://www.example.com",
        help="Enter the business website URL you want to analyze"
    )
    
    # Process button
    if st.button("🚀 Generate Prompt", type="primary", use_container_width=True):
        if not url:
            st.error("❌ Please enter a valid website URL")
            return
        
        # Basic URL validation
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        
        if '.' not in url or len(url) < 10:
            st.error("❌ Please enter a valid website URL (e.g., https://example.com)")
            return
        
        # Initialize session state
        if 'process_complete' not in st.session_state:
            st.session_state.process_complete = False
        if 'final_prompt' not in st.session_state:
            st.session_state.final_prompt = None
        if 'gpt_response' not in st.session_state:
            st.session_state.gpt_response = None
        if 'scraped_data' not in st.session_state:
            st.session_state.scraped_data = None
        
        # Progress tracking
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        try:
            # Step 1: Scrape business data
            status_text.text("🔍 Scraping website data...")
            progress_bar.progress(20)
            
            scraped_data = scrape_business_info_with_ai(url)
            if not scraped_data:
                st.error("❌ Failed to scrape data from the website")
                st.info("💡 Possible reasons:\n• Website might be blocking automated access\n• URL might be incorrect\n• Network connection issues")
                return
            
            st.session_state.scraped_data = scraped_data
            progress_bar.progress(40)
            
            # Step 2: Load template
            status_text.text("📄 Loading prompt template...")
            template = load_template()
            if not template:
                return
            
            progress_bar.progress(60)
            
            # Step 3: Map data to template placeholders
            status_text.text("🔄 Mapping data to template...")
            data_mapping = map_scraped_data_to_template(scraped_data)
            
            # Step 4: Replace placeholders
            status_text.text("✏️ Generating final prompt...")
            final_prompt = replace_placeholders(template, data_mapping)
            st.session_state.final_prompt = final_prompt
            
            progress_bar.progress(80)
            
            # Step 5: Test GPT API if requested
            if test_gpt:
                status_text.text("🤖 Testing with GPT API...")
                test_success, gpt_response, api_stats = test_gpt_response(final_prompt, model)
                st.session_state.gpt_response = gpt_response if test_success else None
                
                if test_success:
                    progress_bar.progress(100)
                    status_text.text("✅ Process completed successfully!")
                else:
                    progress_bar.progress(90)
                    status_text.text("⚠️ Process completed with GPT API error")
            else:
                progress_bar.progress(100)
                status_text.text("✅ Process completed successfully!")
            
            st.session_state.process_complete = True
            
        except Exception as e:
            st.error(f"❌ An error occurred: {str(e)}")
            return
        
        # Clear progress indicators
        time.sleep(1)
        progress_bar.empty()
        status_text.empty()
    
    # Display results if process is complete
    if st.session_state.get('process_complete', False):
        st.markdown("---")
        st.markdown('<h2 class="sub-header">📊 Results</h2>', unsafe_allow_html=True)
        
        # Data summary
        if st.session_state.scraped_data:
            scraped_data = st.session_state.scraped_data
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("**🏢 Basic Information**")
                st.write(f"• Company: {scraped_data.get('company_name', 'Not found')}")
                st.write(f"• Address: {scraped_data.get('address', 'Not found')}")
                st.write(f"• Phone: {scraped_data.get('phone_number', 'Not found')}")
                st.write(f"• Email: {scraped_data.get('email', 'Not found')}")
            
            with col2:
                st.markdown("**🛠️ Services & Pricing**")
                services = scraped_data.get('services_list', 'Not found')
                st.write(f"• Services: {services}")
                pricing = scraped_data.get('pricing', 'Not found')
                st.write(f"• Pricing: {pricing}")
        
        # Final prompt preview
        if st.session_state.final_prompt:
            st.markdown("**📝 Generated Prompt Preview**")
            with st.expander("View Full Prompt", expanded=False):
                st.text_area("Final Prompt", st.session_state.final_prompt, height=300)
            
            # Download options
            st.markdown("**📥 Download Options**")
            
            col1, col2 = st.columns(2)
            
            with col1:
                # TXT download
                txt_content = st.session_state.final_prompt
                st.download_button(
                    label="📄 Download as TXT",
                    data=txt_content,
                    file_name="business_prompt.txt",
                    mime="text/plain",
                    use_container_width=True
                )
            
            with col2:
                # DOCX download
                if st.session_state.gpt_response:
                    company_name = st.session_state.scraped_data.get('company_name', 'Business') if st.session_state.scraped_data else 'Business'
                    docx_bytes = create_docx_download(
                        st.session_state.final_prompt,
                        st.session_state.gpt_response,
                        company_name
                    )
                    
                    st.download_button(
                        label="📄 Download as DOCX",
                        data=docx_bytes.getvalue(),
                        file_name=f"{company_name}_analysis.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                else:
                    st.info("Enable 'Test with GPT API' to download DOCX")
        
        # GPT Response display
        if st.session_state.gpt_response:
            st.markdown("**🤖 AI Analysis & Recommendations**")
            st.markdown(st.session_state.gpt_response)
            
            # API stats
            if hasattr(st.session_state, 'api_stats') and st.session_state.api_stats:
                with st.expander("📊 API Usage Statistics"):
                    stats = st.session_state.api_stats
                    st.write(f"• Total Tokens: {stats['total_tokens']}")
                    st.write(f"• Prompt Tokens: {stats['prompt_tokens']}")
                    st.write(f"• Response Tokens: {stats['completion_tokens']}")

if __name__ == "__main__":
    main() 